from __future__ import annotations
from typing import List, Tuple


class DocxExtractor:
    """Extract text from .docx files, preserving heading structure as Markdown."""

    def extract(self, path: str) -> List[Tuple[int, str]]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Install python-docx: pip install 'docintel[office]'")

        doc = Document(path)
        sections: list[tuple[int, str]] = []
        buffer: list[str] = []
        section_num = 1

        def _flush() -> None:
            nonlocal section_num
            text = "\n".join(buffer).strip()
            if text:
                sections.append((section_num, text))
                section_num += 1
            buffer.clear()

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else ""
            # Convert heading styles to Markdown heading markers
            if style.startswith("Heading"):
                try:
                    level = int(style.split()[-1])
                except ValueError:
                    level = 1
                if level <= 2:
                    _flush()
                text = "#" * min(level, 6) + " " + text
            buffer.append(text)

        _flush()

        # Also pull text from tables
        for table in doc.tables:
            rows: list[str] = []
            for i, row in enumerate(table.rows):
                cells = [c.text.strip() for c in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
            if rows:
                sections.append((section_num, "\n".join(rows)))
                section_num += 1

        return sections if sections else [(1, "")]
